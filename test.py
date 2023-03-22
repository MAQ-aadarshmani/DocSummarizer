# import docx

# # open the docx file
# doc = docx.Document(r'C:\Users\USER-AMAAN\Documents\DocSummarizer\basedonchatgpt\BRD.docx')

# # specify the heading level you want to extract
# heading_level = 1

# # initialize an empty string to hold the text content
# heading_text = ''

# # loop through all the paragraphs in the document
# for para in doc.paragraphs:
#     # check if the paragraph is a heading of the specified level
#     if para.style.name.startswith('Heading ' + str(heading_level)):
#         # add the text of the heading to the string
#         heading_text += para.text + '\n'
#         # loop through all subsequent paragraphs until the next heading of the same or higher level is found
#         for p in para._element.xpath('./following-sibling::*'):
#             # check if the paragraph is a heading of the same or higher level
#             if p.tag.endswith('}p'):
#                 if p.xpath('.//w:pStyle') and p.xpath('.//w:pStyle')[0].attrib.get('w:val') and p.xpath('.//w:pStyle')[0].attrib.get('w:val').startswith('Heading '):
#                     break
#                 # add the text of the paragraph to the string if it is not None
#                 if p.text is not None:
#                     heading_text += p.text + '\n'

# # print the heading text
# print(heading_text)


import docx

# open the docx file
doc = docx.Document(r'C:\Users\USER-AMAAN\Documents\DocSummarizer\basedonchatgpt\BRD.docx')

# initialize an empty string to hold the text content
heading_text = ''

# set a flag to indicate if we are currently inside a Heading 1 section
inside_heading_1 = False
specific_heading = 'Core Value Focus Area – Revenue/Core Growth Strategy'
# loop through all the paragraphs in the document
for para in doc.paragraphs:
    # check if the paragraph is a Heading 1
    if para.style.name.startswith('Heading 1') and inside_heading_1:
        break
    if para.style.name.startswith('Heading 1') and para.text == specific_heading:
        heading_text += para.text + '\n'
        inside_heading_1 = True
    # check if the paragraph is not a Heading 1 and we are currently inside a Heading 1 section
    elif inside_heading_1:
        # check if the paragraph has no style or is bold
        if (not para.style.name.startswith('Heading 1')):
            # add the text of the paragraph to the string if it is not None
            if para.text is not None:
                heading_text += para.text + '\n'
        # if the paragraph has a style starting with "Heading", we are no longer inside a Heading 1 section
        elif para.style.name.startswith('Heading 1'):
            inside_heading_1 = False

# print the content between the Heading 1 sections
print(heading_text)





# import docx

# # open the docx file
# doc = docx.Document(r'C:\Users\USER-AMAAN\Documents\DocSummarizer\basedonchatgpt\BRD.docx')

# # initialize an empty string to hold the text content
# heading_text = ''

# # set a flag to indicate if we are currently inside a Heading 1 section
# inside_heading_1 = False
# # specific_heading = 'CORE VALUE FOCUS AREA - REVENUE/CORE GROWTH STRATEGY'
# # loop through all the paragraphs in the document
# for para in doc.paragraphs:
#     # check if the paragraph is a Heading 1
#     if para.style.name.startswith('Heading 1'):
#         print(para.text)

# # print the content between the Heading 1 sections
# # print(heading_text)
