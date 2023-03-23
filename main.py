import customtkinter as tk
from tkinter import ttk
import openai
from apikey import API_KEY
from tkinter import filedialog
from tkinter import messagebox
import os
import docx
from fpdf import FPDF

pdf = FPDF("P", "mm", "A4")
pdf.add_font("OpenSans", "", os.getcwd() + r"\OpenSans-LightItalic.ttf")
pdf.add_font("OpenSans", "B", os.getcwd() + r"\OpenSans-SemiBoldItalic.ttf")
tk.set_appearance_mode("dark")
tk.set_default_color_theme("dark-blue")


class App:
    def __init__(self, master):
        self.master = master
        self.file_path = ""
        self.summary = ""
        self.file_path_list = []
        self.heading_options = []
        self.desired_content = ""
        self.ext = ""
        self.selected_heading = ""
        self.selected_subheading = ""

        def heading_dropdown_callback(choice):
            self.selected_heading = choice
            self.subheading_dropdown.configure(state="normal")
            self.subheading_dropdown.set("Select a subheading")

            # Subheading values
            self.subheading_dropdown.configure(
                values=subheading_dropdown_values(choice)
            )

            # print("combobox dropdown clicked:", choice)
            doc = docx.Document(self.file_path)
            inside_heading_1 = False
            # loop through all the paragraphs in the document
            for para in doc.paragraphs:
                # check if the paragraph is a Heading 1
                if para.style.name.startswith("Heading 1") and inside_heading_1:
                    break
                if para.style.name.startswith("Heading 1") and para.text == choice:
                    self.desired_content += para.text + "\n"
                    inside_heading_1 = True
                # check if the paragraph is not a Heading 1 and we are currently inside a Heading 1 section
                elif inside_heading_1:
                    if para.text is not None:
                        self.desired_content += para.text + "\n"
            self.desired_content += "\n Summarize it and list out important dimensions, KPI's and metric definitions if any in short."

        def subheading_dropdown_callback(choice):
            self.selected_subheading = choice
            doc = docx.Document(self.file_path)
            inside_heading_1 = False
            inside_subheading_1 = False
            self.desired_content = ""
            # loop through all the paragraphs in the document
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading 2") and inside_subheading_1:
                    break
                # check if the paragraph is a Heading 1
                if (
                    para.style.name == "Heading 1"
                    and para.text == self.selected_heading
                ):
                    inside_heading_1 = True
                if (
                    para.style.name.startswith("Heading 2")
                    and para.text == choice
                    and inside_heading_1
                ):
                    self.desired_content += para.text + "\n"
                    inside_subheading_1 = True
                # check if the paragraph is not a Heading 1 and we are currently inside a Heading 1 section
                elif inside_heading_1 and inside_subheading_1:
                    if para.text is not None:
                        self.desired_content += para.text + "\n"
            self.desired_content += "\n Summarize it and list out important dimensions, KPI's and metric definitions if any in short."

        def subheading_dropdown_values(choice):
            doc = docx.Document(self.file_path)
            # print(doc)
            sub_list = []
            inside_heading1 = False
            for para in doc.paragraphs:
                if para.style.name == "Heading 1" and para.text == choice:
                    inside_heading1 = True
                elif para.style.name == "Heading 2" and inside_heading1:
                    sub_list.append(para.text)
                elif para.style.name == "Heading 1" and inside_heading1:
                    break
            return sub_list

        # create select file button
        self.select_file_button = tk.CTkButton(
            self.master, text="Select File", command=self.select_file
        )
        self.select_file_button.grid(row=0, column=3)

        # Configure 6 rows and 4 columns
        for i in range(16):
            root.rowconfigure(i, weight=1)
        for i in range(7):
            root.columnconfigure(i, weight=1)

        # create label to show selected file path
        self.file_path_label = tk.CTkLabel(self.master, text="No file selected.")
        self.file_path_label.grid(row=1, column=1, columnspan=5, sticky="ew")

        # Create the first Combobox and place it on the left side
        self.heading_dropdown = tk.CTkComboBox(
            self.master,
            values=self.heading_options,
            command=heading_dropdown_callback,
            state="disabled",
        )
        self.heading_dropdown.grid(row=2, column=1, columnspan=2, sticky="ew")

        # Create the second Combobox and place it on the right side
        self.subheading_dropdown = self.subheading_dropdown = tk.CTkComboBox(
            self.master,
            values=self.heading_options,
            command=subheading_dropdown_callback,
            state="disabled",
        )
        self.heading_dropdown.set("Select a Subheading")
        self.subheading_dropdown.grid(row=2, column=4, columnspan=2, sticky="ew")

        # create button to show file content
        self.show_content_button = tk.CTkButton(
            self.master, text="Show Summary", command=self.show_summary
        )
        self.show_content_button.grid(row=3, column=3)

        # create text area to show file content
        self.content_text = tk.CTkTextbox(self.master)
        self.content_text.grid(row=4, column=1, columnspan=5, rowspan=8, sticky="nsew")

        master.title("File Save")

        # Create a button to save output as PDF
        self.pdf_button = tk.CTkButton(
            self.master, text="Save as PDF", command=self.save_as_pdf
        )
        self.pdf_button.grid(row=13, column=1, columnspan=2, sticky="ew")

        # Create a button to save output as TXT
        self.txt_button = tk.CTkButton(
            self.master, text="Save as TXT", command=self.save_as_txt
        )
        self.txt_button.grid(row=13, column=4, columnspan=2, sticky="ew")

    def save_as_pdf(self):
        # Add a page to the PDF object
        pdf.add_page()
        pdf.set_font("OpenSans", "B", 12)
        pdf.cell(
            0,
            10,
            "Generated Summary",
            border=False,
            ln=1,
            align="C",
        )
        pdf.cell(
            0,
            10,
            "for Heading: "
            + self.selected_heading
            + " and Subheading: "
            + self.selected_subheading,
            border=False,
            ln=1,
            align="C",
        )

        pdf.ln(10)
        pdf.set_font("OpenSans", "", size=12)
        pdf.set_fill_color(255, 220, 220)
        text = self.multiline_it(self.summary)

        # Set the cell width and write the text to the cell
        for i in text:
            pdf.cell(0, 10, i, ln=1)

        # Save the PDF file to the downloads folder
        pdf_file = os.path.join(os.path.expanduser("~"), "Downloads", "output.pdf")
        pdf.output(pdf_file)
        messagebox.showinfo("Success", "Successfully saved in "+pdf_file)

    # helper function for pdf creator
    def multiline_it(self, str1):
        str1 = str1.split("\n")
        result = [""]
        chars = 80
        for i in str1:
            str2 = i.split(" ")
            c = 0
            for j in str2:
                c += len(j)
                if c > chars:
                    result[-1] += j
                    result.append("")
                    c = 0
                else:
                    result[-1] += j + " "
            result.append("")
        return result

    # Save the text file to the downloads folder
    def save_as_txt(self):
        txt_file = os.path.join(os.path.expanduser("~"), "Downloads", "output.txt")
        with open(txt_file, "w") as f:
            f.write(
                "Generated Summary for "
                + "Heading: "
                + self.selected_heading
                + " and Subheading: "
                + self.selected_subheading
                + "\n\n"
                + self.summary
            )
        messagebox.showinfo("Success", "Successfully saved in "+txt_file)

    # File selection dialog
    def select_file(self):
        self.file_path = filedialog.askopenfilename()
        self.file_path_label.configure(text=self.file_path)
        self.file_path_list = self.file_path.split(".")
        self.ext = self.file_path_list[-1]
        if self.ext == "docx":
            self.heading_dropdown.configure(state="normal")
            self.heading_dropdown.set("Select a Heading")
            doc = docx.Document(self.file_path)
            for para in doc.paragraphs:
                if para.style.name == "Heading 1":
                    self.heading_options.append(para.text)
            self.heading_dropdown.configure(values=self.heading_options)
        elif self.ext == "pdf":
            pass
        else:
            pass

    # API Call and summary texbox
    def show_summary(self):
        if self.file_path:

            if self.ext != "docx":
                with open(self.file_path, "r") as file:
                    self.desired_content = file.read()
                self.desired_content += "\n Summarize it."

            try:
                openai.api_key = API_KEY
                completion = openai.Completion.create(
                    engine="text-davinci-003",
                    prompt=self.desired_content,
                    max_tokens=300,
                )
                self.summary = completion.choices[0]["text"].strip()
                self.content_text.delete("1.0", tk.END)
                self.content_text.insert(tk.END, self.summary)

            except:
                messagebox.showerror(
                    "Error", "Error: Your file content is exceeding the input limit!!!"
                )
        else:
            self.content_text.delete("1.0", tk.END)
            messagebox.showerror("Error", "Please select a file first.")


root = tk.CTk()
root.title("File Summarizer")
root.geometry("800x500")
app = App(root)
root.mainloop()
